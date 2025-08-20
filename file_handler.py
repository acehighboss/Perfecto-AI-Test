from __future__ import annotations
from typing import Any, Dict, List, Optional, Tuple, Iterable, Union
import io
import os
import re
import csv

# LangChain Document 호환
try:
    from langchain_core.documents import Document
except Exception:
    try:
        from langchain.schema import Document  # type: ignore
    except Exception:
        class Document:  # type: ignore
            def __init__(self, page_content: str, metadata: Optional[Dict[str, Any]] = None):
                self.page_content = page_content
                self.metadata = metadata or {}

# 선택적 의존성
try:
    from pypdf import PdfReader  # pip install pypdf
except Exception:
    try:
        from PyPDF2 import PdfReader  # type: ignore
    except Exception:
        PdfReader = None  # type: ignore

try:
    import docx  # python-docx
except Exception:
    docx = None  # type: ignore

try:
    from bs4 import BeautifulSoup  # pip install beautifulsoup4
except Exception:
    BeautifulSoup = None  # type: ignore

try:
    import chardet  # pip install chardet
except Exception:
    chardet = None  # type: ignore


# -----------------------------------------------------------------------------
# 유틸
# -----------------------------------------------------------------------------
_TEXT_EXTS = {".txt", ".md", ".csv", ".json", ".yaml", ".yml", ".log"}
_HTML_EXTS = {".html", ".htm"}
_SUB_EXTS  = {".srt", ".vtt"}
_DOCX_EXTS = {".docx"}
_PDF_EXTS  = {".pdf"}


def _ext(path_or_name: str) -> str:
    return os.path.splitext(path_or_name.lower())[1]


def _is_filelike(obj: Any) -> bool:
    return hasattr(obj, "read")


def _detect_encoding(b: bytes) -> str:
    if not b:
        return "utf-8"
    if chardet is not None:
        try:
            guess = chardet.detect(b) or {}
            enc = guess.get("encoding")
            if enc:
                return enc
        except Exception:
            pass
    return "utf-8"


def _to_bytes(obj: Union[bytes, bytearray, io.BufferedIOBase, io.BytesIO, Any]) -> bytes:
    if isinstance(obj, (bytes, bytearray)):
        return bytes(obj)
    if _is_filelike(obj):
        pos = None
        try:
            pos = obj.tell()
        except Exception:
            pos = None
        data = obj.read()
        try:
            if pos is not None:
                obj.seek(pos)
        except Exception:
            pass
        if isinstance(data, str):
            return data.encode("utf-8", errors="replace")
        return data or b""
    if isinstance(obj, str) and os.path.exists(obj):
        with open(obj, "rb") as f:
            return f.read()
    return b""


def _bytes_to_text(b: bytes) -> str:
    enc = _detect_encoding(b)
    try:
        return b.decode(enc, errors="replace")
    except Exception:
        return b.decode("utf-8", errors="replace")


def _strip_html(html: str) -> str:
    if not html:
        return ""
    if BeautifulSoup is not None:
        try:
            soup = BeautifulSoup(html, "html.parser")
            for tag in soup(["script", "style", "noscript"]):
                tag.decompose()
            text = soup.get_text("\n")
            # 공백 정리
            text = re.sub(r"[ \t]+", " ", text)
            text = re.sub(r"\n{3,}", "\n\n", text)
            return text.strip()
        except Exception:
            pass
    # 폴백(간단 태그 제거)
    text = re.sub(r"<(script|style)[\s\S]*?</\1>", "", html, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# -----------------------------------------------------------------------------
# 파서: PDF
# -----------------------------------------------------------------------------
def _parse_pdf_to_documents(name: str, raw: bytes) -> List[Document]:
    if PdfReader is None:
        raise ImportError(
            "PDF 처리를 위해 pypdf(또는 PyPDF2)가 필요합니다. `pip install pypdf` 를 설치하세요."
        )
    bio = io.BytesIO(raw)
    reader = PdfReader(bio)
    docs: List[Document] = []
    for i, page in enumerate(reader.pages):
        try:
            txt = page.extract_text() or ""
        except Exception:
            txt = ""
        txt = txt.strip()
        if not txt:
            continue
        meta = {
            "source": name,
            "type": "pdf",
            "page": i + 1,  # 1-based
        }
        docs.append(Document(page_content=txt, metadata=meta))
    return docs or [Document(page_content="", metadata={"source": name, "type": "pdf"})]


# -----------------------------------------------------------------------------
# 파서: DOCX
# -----------------------------------------------------------------------------
def _parse_docx_to_documents(name: str, raw: bytes) -> List[Document]:
    if docx is None:
        # docx 없으면 텍스트로 시도
        return [_as_text_document(name, raw, ftype="docx")]
    try:
        bio = io.BytesIO(raw)
        d = docx.Document(bio)
        paras = [p.text for p in d.paragraphs if p.text and p.text.strip()]
        text = "\n".join(paras).strip()
        meta = {"source": name, "type": "docx"}
        return [Document(page_content=text, metadata=meta)]
    except Exception:
        return [_as_text_document(name, raw, ftype="docx")]


# -----------------------------------------------------------------------------
# 파서: SRT/VTT
# -----------------------------------------------------------------------------
_SRT_TIME = re.compile(r"(\d{1,2}):(\d{2}):(\d{2}),(\d{3})")
_VTT_TIME = re.compile(r"(\d{1,2}):(\d{2}):(\d{2})\.(\d{3})")


def _to_seconds(h: int, m: int, s: int, ms: int) -> float:
    return h * 3600 + m * 60 + s + ms / 1000.0


def _parse_srt(content: str) -> List[Tuple[float, float, str]]:
    lines = content.splitlines()
    cues: List[Tuple[float, float, str]] = []
    buf: List[str] = []
    start = end = 0.0

    def flush():
        nonlocal buf, start, end
        if buf:
            text = "\n".join([x for x in buf if x.strip()]).strip()
            if text:
                cues.append((start, end, text))
        buf = []

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        # 인덱스 줄(숫자) skip
        if re.fullmatch(r"\d+", line):
            i += 1
            if i >= len(lines):
                break
            # 타임라인
            tline = lines[i].strip()
            # 00:00:00,000 --> 00:00:02,000
            parts = re.split(r"\s*-->\s*", tline)
            if len(parts) == 2:
                m1 = _SRT_TIME.search(parts[0])
                m2 = _SRT_TIME.search(parts[1])
                if m1 and m2:
                    start = _to_seconds(*map(int, m1.groups()))
                    end = _to_seconds(*map(int, m2.groups()))
                    i += 1
                    buf = []
                    # 본문
                    while i < len(lines) and lines[i].strip() != "":
                        buf.append(lines[i])
                        i += 1
                    flush()
        i += 1
    return cues


def _parse_vtt(content: str) -> List[Tuple[float, float, str]]:
    # WEBVTT 헤더 제거
    content = re.sub(r"^\ufeff?WEBVTT.*?$", "", content, flags=re.I | re.M).strip()
    cues: List[Tuple[float, float, str]] = []
    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        # 타임라인: 00:00:00.000 --> 00:00:02.000
        if "-->" in line:
            parts = re.split(r"\s*-->\s*", line)
            if len(parts) == 2:
                m1 = _VTT_TIME.search(parts[0])
                m2 = _VTT_TIME.search(parts[1])
                if m1 and m2:
                    start = _to_seconds(*map(int, m1.groups()))
                    end = _to_seconds(*map(int, m2.groups()))
                    i += 1
                    text_lines: List[str] = []
                    while i < len(lines) and lines[i].strip() != "":
                        text_lines.append(lines[i])
                        i += 1
                    text = "\n".join([t for t in text_lines if t.strip()]).strip()
                    if text:
                        cues.append((start, end, text))
        i += 1
    return cues


def _parse_subtitles_to_documents(name: str, raw: bytes, ftype: str) -> List[Document]:
    text = _bytes_to_text(raw)
    cues = _parse_srt(text) if ftype == "srt" else _parse_vtt(text)
    docs: List[Document] = []
    for idx, (start, end, t) in enumerate(cues, 1):
        meta = {
            "source": name,
            "type": ftype,
            "start": float(start),
            "end": float(end),
            "index": idx,
        }
        docs.append(Document(page_content=t, metadata=meta))
    return docs or [Document(page_content="", metadata={"source": name, "type": ftype})]


# -----------------------------------------------------------------------------
# 파서: 텍스트/CSV/HTML 일반
# -----------------------------------------------------------------------------
def _as_text_document(name: str, raw: bytes, ftype: str = "text") -> Document:
    text = _bytes_to_text(raw)
    if ftype in {"html", "htm"}:
        text = _strip_html(text)
    elif ftype == "csv":
        # CSV를 간단히 TSV 형태 문자열로 정규화
        try:
            sio = io.StringIO(text)
            reader = csv.reader(sio)
            rows = ["\t".join(row) for row in reader]
            text = "\n".join(rows)
        except Exception:
            pass
    return Document(page_content=text.strip(), metadata={"source": name, "type": ftype})


# -----------------------------------------------------------------------------
# 공개 함수: get_documents_from_files
# -----------------------------------------------------------------------------
def get_documents_from_files(
    files: Iterable[Union[str, bytes, io.BytesIO, Any]],
    *,
    max_file_size_mb: int = 50,
) -> List[Document]:
    """
    다양한 입력 형태(files)를 LangChain Document 리스트로 변환.
    - files: 스트림릿 UploadedFile, 파일 경로(str), 바이트, 파일-like 등 혼합 가능
    """
    max_bytes = max_file_size_mb * 1024 * 1024
    out: List[Document] = []

    for f in files or []:
        # 이름/경로/확장자 추출
        if _is_filelike(f):
            name = getattr(f, "name", "uploaded_file")
        elif isinstance(f, (bytes, bytearray)):
            name = "bytes_input"
        elif isinstance(f, str):
            name = os.path.basename(f)
        else:
            name = getattr(f, "name", str(f))

        ext = _ext(name)

        # 바이트 로드
        raw = _to_bytes(f)
        if len(raw) > max_bytes:
            # 용량 초과 시 스킵 (필요 시 경고 로그)
            continue

        # 타입 분기
        try:
            if ext in _PDF_EXTS:
                out.extend(_parse_pdf_to_documents(name, raw))
            elif ext in _SUB_EXTS:
                out.extend(_parse_subtitles_to_documents(name, raw, ftype=ext.lstrip(".")))
            elif ext in _DOCX_EXTS:
                out.extend(_parse_docx_to_documents(name, raw))
            elif ext in _HTML_EXTS:
                out.append(_as_text_document(name, raw, ftype="html"))
            elif ext in _TEXT_EXTS:
                ftype = ext.lstrip(".") if ext else "text"
                out.append(_as_text_document(name, raw, ftype=ftype))
            else:
                # 알 수 없는 형식은 텍스트로 시도
                out.append(_as_text_document(name, raw, ftype=ext.lstrip(".") or "bin"))
        except Exception as e:
            # 파싱 실패 시 빈 문서라도 반환(출처 보존)
            out.append(Document(page_content="", metadata={"source": name, "error": str(e)}))

    return out
